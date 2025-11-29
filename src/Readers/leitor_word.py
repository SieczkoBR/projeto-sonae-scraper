import docx
import sqlite3
import os
# Importamos a função de segurança
from Readers.criptograph import encriptar_dado

CAMINHO_ARQUIVO_WORD = "data/relatorio_crm.docx"
CAMINHO_BANCO = "data/projetos_sonae.db"


def ler_word(caminho_arquivo):
    """
    Lê um arquivo Word e retorna todo o texto extraído.
    
    Args:
        caminho_arquivo: Caminho para o arquivo Word (.docx)
        
    Returns:
        String contendo todo o texto do documento
    """
    try:
        documento = docx.Document(caminho_arquivo)
        texto_completo = ""
        
        for paragrafo in documento.paragraphs:
            texto_completo += paragrafo.text + "\n"
        
        return texto_completo
        
    except Exception as e:
        print(f"Erro ao ler Word: {e}")
        return None


def processar_dados_word():
    conexao = None
    try:
        # --- ETAPA 1: LER O ARQUIVO WORD ---
        documento = docx.Document(CAMINHO_ARQUIVO_WORD)
        print(f"Arquivo '{CAMINHO_ARQUIVO_WORD}' lido com sucesso.")

        # --- ETAPA 2: A "PESCARIA" AVANÇADA ---
        print("Iniciando 'modo detetive 3.1'...")
        
        dados_encontrados = {}
        modo_captura = None 
        texto_capturado = [] 

        # Marcadores baseados no seu arquivo real
        MARCADORES = {
            "Nome do Projeto:": "nome_projeto",
            "Responsável:": "responsavel",
            "Status:": "status",
            "Data:": "data_ultima_atualizacao",
            "1. Resumo Executivo": "resumo_executivo",
            "2. Progresso Atual": "progresso_atual",
            "3. Principais Desafios": "principais_desafios",
            "4. Ações Corretivas": "acoes_corretivas",
            "5. Perspectiva": "perspectiva"
        }
        
        for paragrafo in documento.paragraphs:
            texto_linha = paragrafo.text.strip()
            if not texto_linha: continue

            novo_modo = None
            for marcador, chave in MARCADORES.items():
                if texto_linha.startswith(marcador):
                    novo_modo = chave
                    if chave in ["nome_projeto", "responsavel", "status", "data_ultima_atualizacao"]:
                        dados_encontrados[chave] = texto_linha.split(':', 1)[1].strip()
                        novo_modo = None
                    break 

            if novo_modo:
                if modo_captura and texto_capturado:
                    dados_encontrados[modo_captura] = "\n".join(texto_capturado)
                modo_captura = novo_modo 
                texto_capturado = [] 
            elif modo_captura:
                if texto_linha.startswith("-"):
                    texto_capturado.append(f"• {texto_linha[1:].strip()}")
                else:
                    texto_capturado.append(texto_linha)

        if modo_captura and texto_capturado:
            dados_encontrados[modo_captura] = "\n".join(texto_capturado)

        # --- ETAPA 3: VERIFICAÇÃO ---
        if 'nome_projeto' not in dados_encontrados:
            print("ERRO: Não foi possível encontrar o marcador 'Nome do Projeto:'.")
            return 

        print(f"Dados detalhados extraídos para: {dados_encontrados.get('nome_projeto')}")

        # --- ETAPA 4: LÓGICA UPSERT COM SEGURANÇA ---
        conexao = sqlite3.connect(CAMINHO_BANCO)
        cursor = conexao.cursor()

        nome = dados_encontrados.get('nome_projeto')
        
        # AQUI APLICAMOS A SEGURANÇA:
        responsavel_seguro = encriptar_dado(dados_encontrados.get('responsavel'))
        
        dados_completos = {
            "responsavel": responsavel_seguro, # Salva criptografado
            "status": dados_encontrados.get('status'),
            "data_ultima_atualizacao": dados_encontrados.get('data_ultima_atualizacao'),
            "fonte_dados": CAMINHO_ARQUIVO_WORD,
            "resumo_executivo": dados_encontrados.get('resumo_executivo'),
            "progresso_atual": dados_encontrados.get('progresso_atual'),
            "principais_desafios": dados_encontrados.get('principais_desafios'),
            "acoes_corretivas": dados_encontrados.get('acoes_corretivas'),
            "perspectiva": dados_encontrados.get('perspectiva')
        }

        cursor.execute("SELECT id FROM projetos WHERE nome_projeto = ?", (nome,))
        projeto_existente = cursor.fetchone()

        if projeto_existente:
            # UPDATE
            id_projeto = projeto_existente[0]
            sql_update = """
            UPDATE projetos SET 
                responsavel = ?, status = ?, data_ultima_atualizacao = ?, fonte_dados = ?,
                resumo_executivo = ?, progresso_atual = ?, principais_desafios = ?,
                acoes_corretivas = ?, perspectiva = ?
            WHERE id = ?
            """
            valores = list(dados_completos.values()) + [id_projeto]
            cursor.execute(sql_update, valores)
            print(f"Sucesso! Projeto '{nome}' ATUALIZADO (Seguro).")
        else:
            # INSERT
            sql_insert = """
            INSERT INTO projetos (
                nome_projeto, responsavel, status, data_ultima_atualizacao, fonte_dados,
                resumo_executivo, progresso_atual, principais_desafios,
                acoes_corretivas, perspectiva
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """
            valores = [nome] + list(dados_completos.values())
            cursor.execute(sql_insert, valores)
            print(f"Sucesso! Projeto '{nome}' INSERIDO (Seguro).")

        conexao.commit()

    except Exception as e:
        print(f"ERRO inesperado ao processar Word: {e}")
        if conexao: conexao.rollback()
    finally:
        if conexao: conexao.close()

if __name__ == "__main__":
    processar_dados_word()